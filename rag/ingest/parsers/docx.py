"""DOCX document parser."""

import asyncio
from io import BytesIO

from rag.ingest.parsers.base import BaseParser, ParseError
from rag.models.document import Document, DocumentType


class DocxParser(BaseParser):
    """Parser for DOCX documents."""

    @property
    def supported_mimetypes(self) -> set[str]:
        return {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    async def parse(self, content: bytes, filename: str) -> Document:
        """Parse DOCX content into a Document."""
        try:
            loop = asyncio.get_event_loop()
            text, metadata = await loop.run_in_executor(
                None, self._parse_sync, content, filename
            )

            return self._create_document(
                content=text,
                doc_type=DocumentType.DOCX,
                filename=filename,
                raw_content=content,
                title=metadata.get("title"),
                metadata=metadata,
            )
        except ImportError as e:
            raise ParseError(
                f"DOCX parsing requires 'python-docx' package: {e}",
                filename=filename,
            ) from e
        except Exception as e:
            raise ParseError(
                f"Failed to parse DOCX: {e}",
                filename=filename,
            ) from e

    def _parse_sync(self, content: bytes, filename: str) -> tuple[str, dict]:
        """Synchronous DOCX parsing."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            # Try unstructured as fallback
            return self._parse_with_unstructured(content, filename)

        doc = DocxDocument(BytesIO(content))

        # Extract text
        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    texts.append(row_text)

        text = "\n\n".join(texts)

        # Extract metadata
        metadata = {}
        core_props = doc.core_properties

        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created_at"] = core_props.created
        if core_props.modified:
            metadata["updated_at"] = core_props.modified
        if core_props.subject:
            metadata["description"] = core_props.subject
        if core_props.keywords:
            metadata["tags"] = [
                k.strip() for k in core_props.keywords.split(",") if k.strip()
            ]

        return text, metadata

    def _parse_with_unstructured(
        self, content: bytes, filename: str
    ) -> tuple[str, dict]:
        """Fallback parsing using unstructured."""
        try:
            from unstructured.partition.docx import partition_docx
        except ImportError:
            raise ParseError(
                "DOCX parsing requires 'python-docx' or 'unstructured[docx]'",
                filename=filename,
            )

        elements = partition_docx(file=BytesIO(content))

        texts = []
        for element in elements:
            if hasattr(element, "text") and element.text:
                texts.append(element.text)

        text = "\n\n".join(texts)
        metadata = {}

        return text, metadata
